# 创建人：QI-BRO
# 开发时间：2024-08-06  13:52
import os
from io import BytesIO
import PyPDF2
from modelscope import  AutoTokenizer, AutoModel
import torch
from docx import Document
import io

torch.manual_seed(0)
#add/replace your embedding model path
tokenizer_bge = AutoTokenizer.from_pretrained('./ssl_chat/model/bge-large-zh-v1.5')
model_bge = AutoModel.from_pretrained('./ssl_chat/model/bge-large-zh-v1.5')

#参数：本地文件路径 输出的存储文件路径
def get_text_from_file(file_path: str, output_file_path: str) -> bool:
    # 确保文件路径不为空
    if not file_path:  
        return False
  
    import os  
    filename, file_extension = os.path.splitext(os.path.basename(file_path))  
    file_extension = file_extension.lower()[1:]  
  
    try:  
        if file_extension == 'txt':  
            with open(file_path, 'r', encoding='utf-8') as file:  
                file_content = file.read()  
        elif file_extension in ('docx', 'pdf'):  
            with open(file_path, 'rb') as file:  
                file_content = file.read()  
                if file_extension == 'docx':  
                    doc = Document(BytesIO(file_content))  
                    file_content = '\n'.join([para.text for para in doc.paragraphs])  
                elif file_extension == 'pdf':  
                    reader = PyPDF2.PdfFileReader(BytesIO(file_content))  
                    file_content = '\n'.join([page.extract_text() for page in reader.pages if page.extract_text()])  
        elif file_extension == 'json':  
            with open(file_path, 'r', encoding='utf-8') as file:  
                file_content = file.read()  
        else:  
            return False 
  
        with open(output_file_path, 'w', encoding='utf-8') as f:  
            f.write(file_content)  
  
        return True
  
    except Exception as e:  
        print(f"遇到错误:{e}")
        return False

def get_embedding(file_path):
    def process_texts(sentences):
        encoded_input = tokenizer_bge(sentences, padding=True, truncation=True, return_tensors='pt')
        with torch.no_grad():
            model_output = model_bge(**encoded_input)
            sentence_embedding = model_output[0][:, 0]
        return sentence_embedding.tolist()
    try:
        with open(file_path, encoding='utf-8') as f:
            line_number = 0
            text = ""
            mid_str=""
            for line in f:
                mid_str+=str(line)
                
        print(f"mid_str 的长度是 {len(mid_str)}")
        texts = chunk_string(mid_str,200)
    
        res_feature = process_texts(texts)
    
        return res_feature, texts
    except Exception as e:
        print(f"遇到报错 : {e}")


def chunk_string(s, chunk_size):  
    texts = [] 
    if len(s) < 200:
        texts.append(s)
        return texts
    else:
        for i in range(0, len(s), chunk_size):  
            texts.append(s[i:i + chunk_size])  
        return texts  
        






def length_detection(a_str, line_str):
    ori_str = a_str
    a_str = a_str + " " + line_str
    if len(a_str) >= 200:
        return False, line_str
    else:
        return True, a_str






def get_text_from_localfile(input_file_path: str, output_file_path: str) -> bool:
    filename = os.path.basename(input_file_path)
    file_extension = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''

    try:
        with open(input_file_path, 'rb') as file:
            file_content = file.read()

        with open(output_file_path, 'w', encoding='utf-8') as f:
            if file_extension in ('txt',):
                f.write(file_content.decode('utf-8'))
            elif file_extension in ('docx',):
                # 对于.docx文件，使用python-docx库提取文本
                doc = Document(io.BytesIO(file_content))
                f.write('\n'.join([para.text for para in doc.paragraphs]))
            elif file_extension == 'pdf':
                reader = PyPDF2.PdfFileReader(io.BytesIO(file_content))
                f.write('\n'.join([page.extract_text() for page in reader.pages if page.extract_text()]))
            else:
                # 不支持的文件格式
                print(f"Unsupported file format: {file_extension}")
                return False

        return True

    except Exception as e:

        print(f"处理文件上传遇到 {filename}: {e}")
        return False
